#!/usr/bin/env python
"""
Corrected Example:
 - Include has_gps and percentile_score in the select list.
 - Use those columns instead of missing_gps_flag for the asterisk logic.
 - Optionally refine your SQL WHERE clause to match exactly the race(s) you want.
 - Ensure matplotlib uses a backend that can render off-screen (Agg).
"""

import os
import sys
import logging
import datetime
import configparser
import pandas as pd
import numpy as np

# IMPORTANT: set a non-GUI backend so charts can be generated headless
import matplotlib
matplotlib.use('agg')
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches
import io
from psycopg2 import pool, DatabaseError

############################################################
# 1) Setup Logging, Configuration, and DB Pool
############################################################
def setup_logging(script_dir, log_file):
    try:
        with open(log_file, 'w'):
            pass
        logger = logging.getLogger()
        if logger.hasHandlers():
            logger.handlers.clear()
        logger.setLevel(logging.INFO)
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(logging.INFO)
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        logging.info("Logging initialized.")
    except Exception as e:
        print(f"Failed to set up logging: {e}", file=sys.stderr)
        sys.exit(1)

def read_config(script_dir, config_relative_path='../../config.ini'):
    try:
        config = configparser.ConfigParser()
        config_file_path = os.path.join(script_dir, config_relative_path)
        if not os.path.exists(config_file_path):
            raise FileNotFoundError(f"Configuration file '{config_file_path}' does not exist.")
        config.read(config_file_path)
        if 'database' not in config:
            raise KeyError("Missing 'database' section in the config file.")
        return config
    except Exception as e:
        logging.error(f"Error reading configuration file: {e}")
        sys.exit(1)

def get_db_pool(config):
    try:
        db_pool_args = {
            'user': config['database']['user'],
            'host': config['database']['host'],
            'port': config['database']['port'],
            'database': config['database']['dbname']
        }
        password = config['database'].get('password')
        if password:
            db_pool_args['password'] = password
            logging.info("Using password authentication.")
        else:
            logging.info("No password in config; relying on .pgpass or other method.")
        db_pool = pool.SimpleConnectionPool(1, 10, **db_pool_args)
        if db_pool:
            logging.info("Connection pool created successfully.")
        return db_pool
    except DatabaseError as e:
        logging.error(f"Database error creating pool: {e}")
        sys.exit(1)
    except KeyError as e:
        logging.error(f"Missing config key: {e}")
        sys.exit(1)
    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        sys.exit(1)

############################################################
# 2) Softmax and Ensemble Calculation
############################################################
def softmax(arr):
    """Compute a 1D softmax for an array of scores."""
    arr = np.array(arr, dtype=float)
    exp_scores = np.exp(arr - np.nanmax(arr))
    return exp_scores / np.nansum(exp_scores)

def ensemble_predictions(df, model_performance, alpha_blend=0.80):
    """
    Example: Weighted ensemble across several model columns, then
             blend with percentile_score or global_speed_score.
    """
    model_cols = list(model_performance.keys())
    weights = np.array([model_performance[col] for col in model_cols], dtype=float)
    norm_weights = weights / weights.sum()

    # Composite from the CatBoost model predictions:
    df['composite_score'] = df[model_cols].apply(
        lambda row: np.dot(row.values, norm_weights), axis=1
    )

    # Blend composite with percentile_score. 
    # If you still prefer blending with global_speed_score, adapt accordingly:
    df['final_score'] = alpha_blend * df['composite_score'] + \
                        (1 - alpha_blend) * df['percentile_score']

    # Group softmax by race:
    def race_softmax(s):
        return softmax(s.values)

    df['winning_probability'] = df.groupby(
        ['course_cd', 'race_date', 'race_number']
    )['final_score'].transform(lambda s: race_softmax(s))

    return df

############################################################
# 3) Generate Race Charts and Write DOCX
############################################################
def add_race_chart(doc, race_df):
    """
    Generate a horizontal bar chart for a race showing each horse's winning probability,
    then insert the chart into the DOCX document.
    """
    race_df_sorted = race_df.sort_values('winning_probability', ascending=True)
    num_horses = len(race_df_sorted)

    fig, ax = plt.subplots(figsize=(6, max(2, num_horses * 0.3)))
    ax.barh(race_df_sorted['horse_name'], race_df_sorted['winning_probability'] * 100, color='skyblue')
    ax.set_xlabel('Winning Probability (%)')
    ax.set_title('Race Winning Probabilities')
    plt.tight_layout()

    # Save plot to buffer
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)

    # Add to doc
    doc.add_picture(buf, width=Inches(6))
    doc.add_paragraph("Figure: Winning Probability Distribution", style='Intense Quote')

def write_predictions_to_docx(pred_df, output_docx):
    """
    Write predictions to a DOCX file as race cards. Each race card includes:
      - A header with race info
      - A table with columns: (Saddle Cloth, Horse Name*, percentile, composite, WinProb)
      - A horizontal bar chart
      - A page break between races
    """
    # Sort
    pred_df = pred_df.sort_values(['track_name', 'race_date', 'race_number', 'saddle_cloth_number'])

    race_groups = pred_df.groupby(['track_name', 'course_cd', 'race_date', 'race_number'], sort=False)
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for (track_name, course_cd, race_date, race_number), race_df in race_groups:
        # Header
        header_text = f"Race: {course_cd} | {race_date.strftime('%Y-%m-%d')} | Race #{int(race_number)} | Track: {track_name}"
        doc.add_heading(header_text, level=1)
        doc.add_paragraph(f"Track: {track_name}")

        # Table with 5 columns
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Saddle Cloth #"
        hdr_cells[1].text = "Horse Name"
        hdr_cells[2].text = "Percentile"
        hdr_cells[3].text = "Composite"
        hdr_cells[4].text = "Win Prob (%)"

        # If any horse has_gps=0, note it after the table
        # (Or set an asterisk on that horse's name)
        for _, row in race_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['saddle_cloth_number'])

            # Put an asterisk if has_gps=0 (or any other rule you want):
            horse_display = row['horse_name']
            if row['has_gps'] == 0:
                horse_display += " *"
            cells[1].text = horse_display

            # Show percentile
            pct = row['percentile_score']
            if pd.notnull(pct):
                cells[2].text = f"{pct:.3f}"
            else:
                cells[2].text = ""

            # Composite
            comp = row['composite_score']
            if pd.notnull(comp):
                cells[3].text = f"{comp:.2f}"
            else:
                cells[3].text = ""

            # Win prob
            wp = row['winning_probability']
            if pd.notnull(wp):
                cells[4].text = f"{wp*100:.2f}%"
            else:
                cells[4].text = ""

        # Add the chart
        add_race_chart(doc, race_df)

        doc.add_page_break()

    doc.save(output_docx)
    logging.info(f"Predictions saved to DOCX file: {output_docx}")
    return output_docx

############################################################
# 4) Main Function
############################################################
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    log_file = os.path.join(script_dir, "Predictions.log")
    setup_logging(script_dir, log_file)
    config = read_config(script_dir)
    db_pool = get_db_pool(config)

    # --- Connect and load predictions ---
    conn = db_pool.getconn()
    
    # IMPORTANT: columns must include has_gps and percentile_score, if that's what you rely on
    columns_to_select = [
        "course_cd", 
        "race_date", 
        "race_number", 
        "track_name", 
        "saddle_cloth_number",
        "horse_name",
        "percentile_score",   # ADDED
        "has_gps",            # ADDED
        "global_speed_score"  # If you still need it
    ]

    # Also include model prediction columns for the ensemble
    model_columns = [
        "YetiRank_top_4_NDCG_top_4",
        "YetiRank_top_1_NDCG_top_1",
        "QueryRMSE_NDCG_top_2",
        "YetiRank_top_2_NDCG_top_4",
        "YetiRank_top_1_NDCG_top_3",
        "YetiRank_top_2_NDCG_top_3",
        "YetiRank_top_2_NDCG_top_2",
        "QueryRMSE_NDCG_top_4",
        "QueryRMSE_NDCG_top_3",
        "YetiRank_top_4_NDCG_top_3"
    ]
    columns_to_select.extend(model_columns)

    sql_select = ", ".join(columns_to_select)
    
    # -- If you want EXACTLY the same rows as your manual query,
    # -- then replicate that WHERE condition:
    #
    #     course_cd = 'TOP'
    #     AND CAST(race_date AS DATE) = '2025-02-22'
    #     AND race_number = 11
    #
    # For example:
    sql = f"""
    SELECT {sql_select}
    FROM predictions_2025_02_20_1
    WHERE course_cd = 'TOP'
      AND CAST(race_date AS DATE) = '2025-02-22'
      AND race_number = 11
    ORDER BY course_cd, race_date, race_number, saddle_cloth_number
    """

    try:
        pred_df = pd.read_sql(sql, conn)
        logging.info(f"Loaded predictions from DB. Shape: {pred_df.shape}")
    except Exception as e:
        logging.error(f"Error loading predictions: {e}", exc_info=True)
        sys.exit(1)
    finally:
        conn.close()

    # Convert race_date to datetime
    pred_df["race_date"] = pd.to_datetime(pred_df["race_date"], errors="coerce")

    # ---- Example ensemble weighting ----
    model_performance = {
        "YetiRank_top_4_NDCG_top_4": 95.47,
        "YetiRank_top_1_NDCG_top_1": 95.13,
        "QueryRMSE_NDCG_top_2": 91.59,
        "YetiRank_top_2_NDCG_top_4": 91.41,
        "YetiRank_top_1_NDCG_top_3": 91.31,
        "YetiRank_top_2_NDCG_top_3": 90.70,
        "YetiRank_top_2_NDCG_top_2": 90.31,
        "QueryRMSE_NDCG_top_4": 90.27,
        "QueryRMSE_NDCG_top_3": 90.22,
        "YetiRank_top_4_NDCG_top_3": 89.25,
    }
    alpha_blend = 0.80

    # Compute ensemble
    pred_df = ensemble_predictions(pred_df, model_performance, alpha_blend)
    logging.info("Ensemble predictions computed.")

    # --- Write to DOCX ---
    output_docx = os.path.join(script_dir, "final_ensemble_predictions.docx")
    write_predictions_to_docx(pred_df, output_docx)
    logging.info("Predictions document created successfully.")

    print(f"Ensemble predictions complete. Document saved to: {output_docx}")

if __name__ == "__main__":
    main()