import os
import sys
import logging
import psycopg2
import pandas as pd
import numpy as np
from datetime import datetime
import configparser
from docx import Document
from docx.shared import Pt, Inches
import matplotlib.pyplot as plt
from io import BytesIO
from psycopg2 import pool, DatabaseError
from src.data_preprocessing.data_prep1.data_utils import initialize_environment
import re
from catboost import CatBoostClassifier
from sklearn.isotonic import IsotonicRegression

############################################################
# 1) Setup Logging, Configuration, and DB Pool
############################################################
def setup_logging(script_dir, log_file):
    try:
        with open(log_file, 'w'):
            pass
        logger = logging.getLogger()
        if logger.hasHandlers():
            logger.handlers.clear()
        logger.setLevel(logging.INFO)
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(logging.INFO)
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        logging.info("Logging initialized.")
    except Exception as e:
        print(f"Failed to set up logging: {e}", file=sys.stderr)
        sys.exit(1)

def read_config(script_dir, config_relative_path='../../config.ini'):
    try:
        config = configparser.ConfigParser()
        config_file_path = os.path.join(script_dir, config_relative_path)
        if not os.path.exists(config_file_path):
            raise FileNotFoundError(f"Configuration file '{config_file_path}' does not exist.")
        config.read(config_file_path)
        if 'database' not in config:
            raise KeyError("Missing 'database' section in the config file.")
        return config
    except Exception as e:
        logging.error(f"Error reading configuration file: {e}")
        sys.exit(1)

def get_db_pool(config):
    try:
        db_pool_args = {
            'user': config['database']['user'],
            'host': config['database']['host'],
            'port': config['database']['port'],
            'database': config['database']['dbname']
        }
        password = config['database'].get('password')
        if password:
            db_pool_args['password'] = password
            logging.info("Using password authentication.")
        else:
            logging.info("No password in config; relying on .pgpass or other method.")
        db_pool = pool.SimpleConnectionPool(1, 10, **db_pool_args)
        if db_pool:
            logging.info("Connection pool created successfully.")
        return db_pool
    except DatabaseError as e:
        logging.error(f"Database error creating pool: {e}")
        sys.exit(1)
    except KeyError as e:
        logging.error(f"Missing config key: {e}")
        sys.exit(1)
    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        sys.exit(1)

############################################################
# 2) Softmax Calculation
############################################################
def softmax(arr):
    """Compute a 1D softmax for an array of scores."""
    arr = np.array(arr, dtype=float)
    exp_scores = np.exp(arr - np.max(arr))
    return exp_scores / np.sum(exp_scores)

############################################################
# 3) Generate and Insert Race Charts into DOCX
############################################################
def add_race_chart(doc, race_df):
    """
    Generate a horizontal bar chart for a race showing each horse's winning probability,
    then insert the chart into the DOCX document.
    """
    # Sort the race DataFrame so that the highest probability appears at the top
    race_df_sorted = race_df.sort_values('winning_probability', ascending=True)
    
    # Create the plot
    num_horses = len(race_df_sorted)
    fig, ax = plt.subplots(figsize=(6, max(2, num_horses * 0.3)))
    ax.barh(race_df_sorted['horse_name'], race_df_sorted['winning_probability'] * 100, color='skyblue')
    ax.set_xlabel('Winning Probability (%)')
    ax.set_title('Race Winning Probabilities')
    plt.tight_layout()
    
    # Save the plot to a BytesIO buffer
    buf = BytesIO()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)
    
    # Insert the image into the DOCX document
    doc.add_picture(buf, width=Inches(6))
    # Optionally, add a caption or spacing
    doc.add_paragraph("Figure: Winning Probability Distribution", style='Intense Quote')
    
    # Close the buffer
    buf.close()

############################################################
# 4) Write Predictions to DOCX
############################################################
def write_predictions_to_docx(pred_df, output_docx):
    """
    Write predictions to a DOCX file as race cards. Each race card includes:
      - Race header information.
      - A table with Saddle Cloth #, Horse Name (asterisk if missing data),
        Global Speed, Composite Score, and Winning Probability.
      - A horizontal bar chart showing the winning probabilities.
      - A note if any horse in the race has missing data.
      - A page break between races.
    """
    # Sort by track, race_date, race_number, and saddle_cloth_number.
    pred_df = pred_df.sort_values(['track_name', 'race_date', 'race_number', 'saddle_cloth_number'])
    # Group by race (assumed unique by track_name, course_cd, race_date, race_number).
    race_groups = pred_df.groupby(['track_name', 'course_cd', 'race_date', 'race_number'], sort=False)
    
    doc = Document()
    # Set default style for the document.
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    for (track_name, course_cd, race_date, race_number), race_df in race_groups:
        logging.info(f"Processing race: {course_cd} | {race_date} | Race #{race_number} | Track: {track_name}")
        
        # Race header.
        header_text = (f"Race: {course_cd} | {race_date.strftime('%Y-%m-%d')} | Race #{int(race_number)} | Track: {track_name}")
        doc.add_heading(header_text, level=1)
        doc.add_paragraph(f"Track: {track_name}")
        
        # Create table with 5 columns.
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Saddle Cloth #"
        hdr_cells[1].text = "Horse Name"
        hdr_cells[2].text = "Global Speed"
        hdr_cells[3].text = "Model Prediction"
        hdr_cells[4].text = "Win Prob (%)"
        
        # Check if any horse has missing data.
        race_has_missing = (race_df['has_gps'].astype(int).sum() > 0)
        
        # Sort by saddle_cloth_number using natural sort.
        # Step 1: Create a new column for sorting
        race_df['saddle_cloth_sortkey'] = race_df['saddle_cloth_number'].apply(natural_sort_key)

        # Step 2: Sort by that new column
        race_df = race_df.sort_values('saddle_cloth_sortkey').reset_index(drop=True)

        # Step 3: Drop the sort key if you don’t need it
        race_df.drop(columns=['saddle_cloth_sortkey'], inplace=True)

        
        # Populate table rows.
        for _, row in race_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['saddle_cloth_number'])
            horse_name = row['horse_name']
            if row.get('has_gps', 0) == 0:
                horse_name += " *"
            cells[1].text = horse_name
            cells[2].text = f"{row['global_speed_score_iq']:.1f}"
            cells[3].text = f"{row['YetiRank_NDCG_top_4']:.2f}"
            cells[4].text = f"{row['winning_probability']*100:.2f}"
        
        # Add the race chart.
        add_race_chart(doc, race_df)
        
        # Add a note if there is missing data.
        if race_has_missing:
            doc.add_paragraph("NOTE: THIS RACE HAS HORSES THAT ARE MISSING CRITICAL DATA ELEMENTS", style='Intense Quote')
        
        # Insert a page break after each race card.
        doc.add_page_break()
    
    doc.save(output_docx)
    logging.info(f"Predictions saved to DOCX file: {output_docx}")
    return output_docx

def natural_sort_key(s):
    """Sort strings containing numbers in a natural way."""
    return [int(text) if text.isdigit() else text.lower() for text in re.split('([0-9]+)', s)]

############################################################
# 5) Main Function
############################################################
def main():
    # Setup environment
    script_dir = os.path.dirname(os.path.abspath(__file__))
    log_file = os.path.join(script_dir, "Predictions.log")
    setup_logging(script_dir, log_file)
    config = read_config(script_dir)
    db_pool = get_db_pool(config)
    spark, jdbc_url, jdbc_properties, parquet_dir, _ = initialize_environment()
    
    # Connect to the database and load predictions
    conn = db_pool.getconn()
    cursor = conn.cursor()
    
    columns_to_select = [
        "course_cd", "horse_id", "group_id", "race_date", "race_number", "post_time", "track_name", "has_gps", "saddle_cloth_number",
        "horse_name", "global_speed_score_iq", "YetiRank_NDCG_top_4"]
    
    sql_select = ", ".join(columns_to_select)
    sql = f"""
    SELECT {sql_select}
    FROM predictions_2025_03_02_1
    WHERE race_date >= CURRENT_DATE - 3
    ORDER BY course_cd, race_date, race_number, saddle_cloth_number
    """
    
    logging.info(f"Executing SQL query: {sql}")

    try:
        cursor.execute(sql)
        pred_df = pd.DataFrame(cursor.fetchall(), columns=columns_to_select)
        logging.info(f"Loaded predictions from DB. Shape: {pred_df.shape}")
        logging.info(f"Predictions DataFrame head: {pred_df.head()}")
    except Exception as e:
        logging.error(f"Error loading predictions: {e}", exc_info=True)
        sys.exit(1)
    finally:
        cursor.close()
        conn.close()
    
    if pred_df.empty:
        logging.error("Predictions DataFrame is empty. Exiting.")
        sys.exit(1)
    
    # Calculate winning probabilities using softmax within each race group.
    def race_softmax(s):
        return softmax(s.values)
    pred_df['winning_probability'] = pred_df.groupby(['course_cd', 'race_date', 'race_number'])['YetiRank_NDCG_top_4'] \
                                        .transform(lambda s: race_softmax(s))
    
    logging.info("Winning probabilities computed.")
    
    # --- Calibration with Isotonic Regression ---
    try:
        conn = db_pool.getconn()
        query = """
        SELECT
            course_cd,
            horse_id,
            race_date,
            race_number,
            prediction AS model_score,
            CASE WHEN official_fin = 1 THEN 1 ELSE 0 END AS won
        FROM catboost_enriched_results
        WHERE data_flag = 'historical'
        AND model_key = 'YetiRank:top=3_NDCG:top=2_20250304_184142'
        """
        cursor = conn.cursor()
        cursor.execute(query)
        df_calibration = pd.DataFrame(cursor.fetchall(), columns=["course_cd", "horse_id", "race_date", "race_number", "model_score", "won"])
        logging.info(f"Loaded calibration data from DB. Shape: {df_calibration.shape}")
    except Exception as e:
        logging.error(f"Error loading calibration data: {e}", exc_info=True)
        sys.exit(1)
    finally:
        cursor.close()
        conn.close()

    # Apply softmax to the calibration data
    df_calibration['raw_prob'] = df_calibration.groupby(['course_cd', 'race_date', 'race_number'])['model_score'] \
                                                .transform(lambda s: softmax(s.values))

    # Fit Isotonic Regression model
    iso = IsotonicRegression(out_of_bounds='clip')
    iso.fit(df_calibration['raw_prob'], df_calibration['won'])
    logging.info("Isotonic Regression model fitted.")

    # Apply the calibration to future predictions
    try:
        conn = db_pool.getconn()
        query = """
        SELECT horse_id, yetirank_ndcg_top_4 AS model_score
        FROM predictions_2025_03_02_1
        """
        cursor = conn.cursor()
        cursor.execute(query)
        df_future = pd.DataFrame(cursor.fetchall(), columns=["horse_id", "model_score"])
        logging.info(f"Loaded future predictions from DB. Shape: {df_future.shape}")
    except Exception as e:
        logging.error(f"Error loading future predictions: {e}", exc_info=True)
        sys.exit(1)
    finally:
        cursor.close()
        conn.close()

    # Apply softmax to the future predictions
    df_future['raw_prob'] = df_future.groupby('horse_id')['model_score'] \
                                     .transform(lambda s: softmax(s.values))

    # Apply the Isotonic Regression model to calibrate the probabilities
    df_future['calibrated_prob'] = iso.predict(df_future['raw_prob'])
    logging.info("Calibrated probabilities computed.")

    # --- Write Predictions to DOCX with Charts ---
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_docx = os.path.join(script_dir, f"final_predictions_CoPilot_{current_date}.docx")
    logging.info(f"Writing predictions to DOCX: {output_docx}")
    write_predictions_to_docx(pred_df, output_docx)
    logging.info("Predictions document created successfully.")
    
    print(f"Predictions complete. Document saved to: {output_docx}")

if __name__ == "__main__":
    main()