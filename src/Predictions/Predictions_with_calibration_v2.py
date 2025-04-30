import os
import sys
import logging
import psycopg2
import pandas as pd
import numpy as np
import re
from datetime import datetime
import configparser
from docx import Document
from docx.shared import Pt, Inches
import matplotlib.pyplot as plt
from io import BytesIO
from psycopg2 import pool, DatabaseError
from src.data_preprocessing.data_prep1.data_utils import initialize_environment
from fractions import Fraction

# For EXACT matches to your distinct probabilities,
# we return typical "nice" fractional odds:
MORN_ODDS_LOOKUP = {
    0.00: "∞",       # 0% => infinite odds
    0.01: "99-1",
    0.02: "50-1",
    0.03: "30-1",
    0.04: "25-1",
    0.05: "20-1",
    0.06: "15-1",
    0.08: "12-1",
    0.09: "10-1",
    0.10: "9-1",
    0.11: "8-1",
    0.13: "7-1",
    0.14: "6-1",
    0.17: "5-1",
    0.18: "9-2",    # ~18.2% is 1/(4.5+1)=.18
    0.20: "4-1",
    0.22: "7-2",    # ~22.2%
    0.25: "3-1",
    0.29: "5-2",    # ~28.6%
    0.33: "2-1",    # 33.3%
    0.36: "9-5",    # ~35.7%
    0.38: "13-8",   # ~38.1%
    0.40: "3-2",    # 40%
    0.42: "7-5",    # ~41.7%
    0.45: "6-5",    # ~45.5%
    0.50: "1-1",    # 50%
    0.56: "4-5",    # ~55.6%
    0.63: "5-8",    # ~61.5% close to 63%
    0.67: "1-2",    # ~66.7%
    0.71: "2-5",    # ~71.4%
    0.83: "1-5",    # ~83.3%
    0.90: "1-9"     # 90%
}
############################################################
# 1) Setup Logging, Config, DB Pool
############################################################
def setup_logging(script_dir, log_file):
    try:
        with open(log_file, 'w'):
            pass
        logger = logging.getLogger()
        if logger.hasHandlers():
            logger.handlers.clear()
        logger.setLevel(logging.INFO)
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(logging.INFO)
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        logging.info("Logging initialized.")
    except Exception as e:
        print(f"Failed to set up logging: {e}", file=sys.stderr)
        sys.exit(1)

def read_config(script_dir, config_relative_path='../../config.ini'):
    try:
        config = configparser.ConfigParser()
        config_file_path = os.path.join(script_dir, config_relative_path)
        if not os.path.exists(config_file_path):
            raise FileNotFoundError(f"Configuration file '{config_file_path}' does not exist.")
        config.read(config_file_path)
        if 'database' not in config:
            raise KeyError("Missing 'database' section in the config file.")
        return config
    except Exception as e:
        logging.error(f"Error reading configuration file: {e}")
        sys.exit(1)

def get_db_pool(config):
    try:
        db_pool_args = {
            'user': config['database']['user'],
            'host': config['database']['host'],
            'port': config['database']['port'],
            'database': config['database']['dbname']
        }
        password = config['database'].get('password')
        if password:
            db_pool_args['password'] = password
            logging.info("Using password authentication.")
        else:
            logging.info("No password in config; relying on .pgpass or other method.")
        db_pool = pool.SimpleConnectionPool(1, 10, **db_pool_args)
        if db_pool:
            logging.info("Connection pool created successfully.")
        return db_pool
    except DatabaseError as e:
        logging.error(f"Database error creating pool: {e}")
        sys.exit(1)
    except KeyError as e:
        logging.error(f"Missing config key: {e}")
        sys.exit(1)
    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        sys.exit(1)

############################################################
# 2) NATURAL SORT HELPER FOR SADDLE_CLOTH
############################################################
def parse_saddle_cloth(s):
    """
    Parse a string like "1", "1A", "3", "3A", "10", "10A" into a tuple (number_part, letter_part).
    This allows us to sort naturally so that 1 < 1A < 2 < 3 < 3A < 10 < 10A.
    """
    match = re.match(r'^(\d+)([A-Za-z]*)$', s)
    if match:
        number_part = int(match.group(1))       # e.g. 3
        letter_part = match.group(2) or ""      # e.g. "A" or ""
        return (number_part, letter_part)
    else:
        return (999999, s)  # fallback if not matching

############################################################
# 3) Charts for the DOCX
############################################################
def add_race_chart(doc, race_df):
    """
    Chart #1: Winning Probability bar chart (sorted ascending).
    """
    race_df_sorted = race_df.sort_values('winning_probability', ascending=True)
    num_horses = len(race_df_sorted)

    fig, ax = plt.subplots(figsize=(8, max(2, num_horses * 0.3)))
    ax.barh(race_df_sorted['horse_name'], race_df_sorted['winning_probability']*100, color='skyblue')
    ax.set_xlabel('Winning Probability (%)')
    ax.set_title('Race Winning Probabilities')
    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)

    doc.add_picture(buf, width=Inches(6))
    buf.close()

# def add_overlay_chart(doc: Document, race_df: pd.DataFrame) -> None:
#     """
#     Create a bar chart showing overlay = (our prob) - (track implied prob).
#     Negative bars mean the horse is overbet by the track (less value).
#     Positive bars mean potentially good value.
#     """
#     if 'morn_odds' not in race_df.columns:
#         doc.add_paragraph("No 'morn_odds' column found, skipping overlay chart.")
#         return

#     # Make a copy so we don't alter the original DataFrame
#     chart_df = race_df.copy()

#     # Convert fractional odds => implied probability: 1 / (odds + 1)
#     chart_df['track_implied_prob'] = 1.0 / (chart_df['morn_odds'].astype(float) + 1.0)

#     chart_df['overlay'] = chart_df['winning_probability'] - chart_df['track_implied_prob']
#     chart_df.sort_values('overlay', ascending=True, inplace=True)

#     num_horses = len(chart_df)

#     fig, ax = plt.subplots(figsize=(8, max(2, num_horses * 0.3)))
#     ax.barh(chart_df['horse_name'], chart_df['overlay'], color='skyblue')
#     ax.set_xlabel('Overlay (Model Prob - Track Prob)')
#     ax.set_title('Value/Overlay Chart')
#     ax.axvline(0, color='black', linewidth=1)  # Center line

#     plt.tight_layout()
#     buf = BytesIO()
#     plt.savefig(buf, format='png')
#     plt.close(fig)
#     buf.seek(0)

#     doc.add_picture(buf, width=Inches(6))
#     buf.close()
        
# def add_overlay_chart(doc, race_df):
#     """
#     Chart #2: Overlay (Value) Chart. Compares our model's probability vs. track implied probability 
#               from fractional morning line odds (in column: morn_odds).
#     """
#     if 'morn_odds' not in race_df.columns:
#         doc.add_paragraph("No 'morn_odds' column found, skipping overlay chart.")
#         return

#     # Make a copy so we don't alter the original DataFrame
#     race_df = race_df.copy()

#     # Convert fractional odds => implied probability
#     # e.g. if morn_odds=5 => 5-1 => implied prob = 1/6=0.1667
#     #    if morn_odds=9/2 => 4.5 => implied prob=1/5.5=0.1818
#     race_df['track_implied_prob'] = 1.0 / (race_df['morn_odds'].astype(float) + 1.0)

#     # overlay = (our prob) - (track implied prob)
#     race_df['overlay'] = race_df['winning_probability'] - race_df['track_implied_prob']

#     # Sort ascending by overlay
#     race_df_sorted = race_df.sort_values('overlay', ascending=True)
#     num_horses = len(race_df_sorted)

#     fig, ax = plt.subplots(figsize=(8, max(2, num_horses*0.3)))
#     ax.barh(race_df_sorted['horse_name'], race_df_sorted['overlay'], color='skyblue')
#     ax.set_xlabel('Overlay Amount (Our Probability - Track Implied)')
#     ax.set_title('Value / Overlay Chart')
#     # Vertical line at 0 for visual reference
#     ax.axvline(0, color='black', linewidth=1)

#     plt.tight_layout()

#     buf = BytesIO()
#     plt.savefig(buf, format='png')
#     plt.close(fig)
#     buf.seek(0)

#     doc.add_picture(buf, width=Inches(6))
#     buf.close()

############################################################
# 4) Write Predictions to DOCX
############################################################
def write_predictions_to_docx(pred_df, bandit_picks_dict, output_docx):
    """
    Write predictions to a DOCX file as race cards, each with:
      - Race header
      - Table of horses
      - Chart #1: Race Probabilities
      - Chart #2: Overlay/Value Chart
      - Exacta advice paragraph
      - Bandit recommendation paragraph
      - Page break
    At the end, adds Multi‐Race Bet Recommendations (DD, Pick 3, Pick 4).
    """
    if pred_df.empty:
        logging.info("No predictions to write to DOCX (DataFrame is empty).")
        return

    # thresholds for gap analysis (tune as needed)
    THR_GAP12 = 0.10
    THR_GAP23 = 0.05

    # Natural sorting by 'saddle_cloth_number'
    pred_df['saddle_cloth_key'] = pred_df['saddle_cloth_number'].apply(parse_saddle_cloth)

    # Sort by track_name, course_cd, race_date, race_number, then natural cloth
    pred_df = pred_df.sort_values(
        ['track_name', 'course_cd', 'race_date', 'race_number', 'saddle_cloth_key']
    )

    race_groups = pred_df.groupby(
        ['track_name', 'course_cd', 'race_date', 'race_number', 'post_time'],
        sort=False
    )

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Collect for multi‐race logic
    multi = []

    for (track_name, course_cd, race_date, race_number, post_time), race_df in race_groups:
        # build race key
        race_key = f"{course_cd}_{race_date.strftime('%Y-%m-%d')}_{int(race_number)}"

        # header
        header_text = (
            f"Race: {course_cd} | {race_date.strftime('%Y-%m-%d')} "
            f"| Race #{int(race_number)} | Post: {post_time} | Track: {track_name}"
        )
        doc.add_heading(header_text, level=1)
        doc.add_paragraph(f"Track: {track_name}")

        # table
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Program #"
        hdr_cells[1].text = "Horse Name"
        hdr_cells[2].text = "Morn Odds"
        hdr_cells[3].text = "Rank"
        hdr_cells[4].text = "Model Score"
        hdr_cells[5].text = "Win Prob (%)"

        race_df = race_df.reset_index(drop=True)
        for _, row in race_df.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row['saddle_cloth_number'])
            cells[1].text = row['horse_name']

            # Morning Odds
            morn_prob = row['morn_odds']
            cells[2].text = MORN_ODDS_LOOKUP.get(morn_prob, "N/A") if pd.notna(morn_prob) else "N/A"

            # Global Speed
            gs = row['rank']
            cells[3].text = f"{gs:.1f}" if pd.notna(gs) else "N/A"

            # Model Score
            ms = row['score']
            cells[4].text = f"{ms:.2f}" if pd.notna(ms) else "N/A"

            # Win Prob
            wp = row['winning_probability']
            cells[5].text = f"{wp * 100:.2f}" if pd.notna(wp) else "N/A"

        # charts
        add_race_chart(doc, race_df)
        #add_overlay_chart(doc, race_df)

        # gap analysis on CatBoost score
        sorted_df = race_df.sort_values('score', ascending=False).reset_index(drop=True)
        s = sorted_df['score'].values
        gap12 = s[0] - s[1] if len(s) > 1 else 0.0
        gap23 = s[1] - s[2] if len(s) > 2 else 0.0

        # exacta advice
        if gap12 > THR_GAP12 and gap23 > THR_GAP23:
            advice = (
                f"Exacta advice: key {sorted_df.loc[0,'saddle_cloth_number']} → "
                f"{sorted_df.loc[1,'saddle_cloth_number']} (1 ticket)"
            )
            confident = True
        elif gap12 > THR_GAP12:
            advice = (
                f"Exacta advice: key {sorted_df.loc[0,'saddle_cloth_number']} → "
                f"[{sorted_df.loc[1,'saddle_cloth_number']},"
                f"{sorted_df.loc[2,'saddle_cloth_number']}] (2 tickets)"
            )
            confident = True
        else:
            advice = (
                f"Exacta advice: box "
                f"[{sorted_df.loc[0,'saddle_cloth_number']},"
                f"{sorted_df.loc[1,'saddle_cloth_number']},"
                f"{sorted_df.loc[2,'saddle_cloth_number']}] (6 tickets)"
            )
            confident = False

        doc.add_paragraph(advice, style='Normal')

        # --- Define Bandit Lookup Key ---
        # Ensure lookup key types match dictionary key types (str, date, int)
        lookup_date = race_date.date() if hasattr(race_date, 'date') else race_date # Convert timestamp/datetime to date
        try:
            lookup_race_num = int(race_number) # Convert float race_number to int
        except (ValueError, TypeError):
            lookup_race_num = -1 # Handle potential conversion errors gracefully
        bandit_lookup_key = (course_cd, lookup_date, lookup_race_num)
        bandit_rec = bandit_picks_dict.get(bandit_lookup_key)
        if bandit_rec:
            bandit_advice = (
                f"Bandit Recommendation: {bandit_rec['chosen_arm']} "
                f"(Expected Net: {bandit_rec['expected_net']:.2f})"
            )
            doc.add_paragraph(bandit_advice, style='Intense Quote') # Use a distinct style if desired
        else:
            # Optionally add a placeholder if no bandit recommendation exists
            # doc.add_paragraph("No bandit recommendation available.", style='Normal') # Uncomment if you want a placeholder
            pass # Or just do nothing if no recommendation

        doc.add_page_break()

        multi.append({
            "race_key":  race_key,
            "track":     track_name,
            "race_num":  race_number,
            "confident": confident
        })

    # remove helper column
    pred_df.drop(columns=['saddle_cloth_key'], inplace=True)

    # Multi‐Race Bet Recommendations
    doc.add_heading("Multi-Race Bet Recommendations", level=1)
    # Daily Double
    for i in range(len(multi) - 1):
        if (multi[i]["confident"] and multi[i+1]["confident"]
            and multi[i]["track"] == multi[i+1]["track"]):
            doc.add_paragraph(
                f"Daily Double advised: {multi[i]['race_key']} + {multi[i+1]['race_key']}"
            )
    # Pick 3
    for i in range(len(multi) - 2):
        if all(multi[j]["confident"] for j in (i, i+1, i+2)) \
           and len({multi[j]["track"] for j in (i, i+1, i+2)}) == 1:
            doc.add_paragraph(
                f"Pick 3 advised: {multi[i]['race_key']}, "
                f"{multi[i+1]['race_key']}, {multi[i+2]['race_key']}"
            )
    # Pick 4
    for i in range(len(multi) - 3):
        if all(multi[j]["confident"] for j in (i, i+1, i+2, i+3)) \
           and len({multi[j]["track"] for j in (i,i+1,i+2,i+3)}) == 1:
            doc.add_paragraph(
                f"Pick 4 advised: {multi[i]['race_key']}, "
                f"{multi[i+1]['race_key']}, {multi[i+2]['race_key']}, {multi[i+3]['race_key']}"
            )

    doc.save(output_docx)
    logging.info(f"Predictions & betting advice saved to {output_docx}")
    
############################################################
# 5) Main
############################################################
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    log_file = os.path.join(script_dir, "Predictions.log")
    setup_logging(script_dir, log_file)
    config = read_config(script_dir)
    db_pool = get_db_pool(config)
    spark, jdbc_url, jdbc_properties, parquet_dir, _ = initialize_environment()

    # ==============================
    # PATH TO BANDIT PICKS CSV
    # ==============================
    bandit_picks_csv_path = "/home/exx/myCode/horse-racing/FoxRiverAIRacing/multi_arm_bandit_future_picks20250429.csv"
    # A) SELECT FROM THE CALIBRATED TABLE
    # ==============================
    conn = db_pool.getconn()
    cursor = conn.cursor()

    # Make sure the column name in your DB is 'morn_odds'
    # in predictions_2025_03_07_1_calibrated
    sql_select = """
        SELECT
            morn_odds,
            group_id,
            TO_CHAR( (post_time)::timestamp, 'HH24:MI' ) AS post_time,
            track_name,
            has_gps,
            horse_name,
            course_cd,
            race_date,
            race_number,
            rank,
            horse_id,
            saddle_cloth_number,
            score,
            calibrated_prob AS winning_probability
        FROM predictions_20250429_221506_1_calibrated
        WHERE race_date >= CURRENT_DATE -- INTERVAL '1 day'
        ORDER BY course_cd, race_date, race_number, saddle_cloth_number
    """

    pred_df = pd.DataFrame()
    try:
        cursor.execute(sql_select)
        rows = cursor.fetchall()
        columns = [
            "morn_odds", "group_id", "post_time", "track_name", "has_gps", "horse_name",
            "course_cd", "race_date", "race_number", "rank",
            "horse_id", "saddle_cloth_number", "score",
            "winning_probability"
        ]
        pred_df = pd.DataFrame(rows, columns=columns)
        logging.info(f"Loaded {len(pred_df)} calibrated predictions from DB.")
    except Exception as e:
        logging.error(f"Error loading calibrated predictions: {e}", exc_info=True)
        sys.exit(1)
    finally:
        cursor.close()
        db_pool.putconn(conn)

    if pred_df.empty:
        logging.info("No rows in the calibrated table for today's date. Exiting.")
        empty_docx = os.path.join(script_dir, "final_predictions_calibrated_empty.docx")
        doc = Document()
        doc.add_heading("No Races Found for Today", level=1)
        doc.save(empty_docx)
        logging.info(f"Saved empty doc: {empty_docx}")
        sys.exit(0)

    # ==============================
    # C) Load Bandit Picks and create lookup dictionary
    # ==============================
    bandit_picks_dict = {}
    try:
        bandit_picks_df = pd.read_csv(bandit_picks_csv_path)
        # Convert race_date string to datetime.date objects to match pred_df
        bandit_picks_df['race_date'] = pd.to_datetime(bandit_picks_df['race_date']).dt.date
        # Convert race_number to int if it's not already
        bandit_picks_df['race_number'] = bandit_picks_df['race_number'].astype(int)

        # Create the dictionary for quick lookup
        for _, row in bandit_picks_df.iterrows():
            key = (row['course_cd'], row['race_date'], row['race_number'])
            bandit_picks_dict[key] = {
                'chosen_arm': row['chosen_arm'],
                'expected_net': row['expected_net']
            }
        logging.info(f"Loaded {len(bandit_picks_df)} bandit picks from {bandit_picks_csv_path}")
    except FileNotFoundError:
        logging.warning(f"Bandit picks file not found at {bandit_picks_csv_path}. Bandit recommendations will be skipped.")
    except Exception as e:
        logging.error(f"Error loading or processing bandit picks CSV: {e}", exc_info=True)
        # Continue without bandit picks if loading fails

    # ==============================
    # B) Generate the DOCX with 2 charts
    # ==============================
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_docx = os.path.join(script_dir, f"final_predictions_calibrated_{current_date}v2.docx")

    write_predictions_to_docx(pred_df, bandit_picks_dict, output_docx)
    logging.info("Predictions document created successfully.")
    print(f"Predictions complete. Document saved to: {output_docx}")


if __name__ == "__main__":
    main()