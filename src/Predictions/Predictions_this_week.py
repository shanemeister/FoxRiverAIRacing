import os
import sys
import logging
import psycopg2
import pandas as pd
import numpy as np
from datetime import datetime
import configparser
from psycopg2 import pool, DatabaseError
from docx import Document
from docx.shared import Inches

from src.data_preprocessing.data_prep1.data_utils import initialize_environment

def setup_logging(script_dir, log_file):
    try:
        with open(log_file, 'w'):
            pass
        logger = logging.getLogger()
        if logger.hasHandlers():
            logger.handlers.clear()
        logger.setLevel(logging.INFO)
        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(logging.INFO)
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)
        logging.info("Logging initialized.")
    except Exception as e:
        print(f"Failed to set up logging: {e}", file=sys.stderr)
        sys.exit(1)

def read_config(script_dir, config_relative_path='../../config.ini'):
    try:
        config = configparser.ConfigParser()
        config_file_path = os.path.join(script_dir, config_relative_path)
        if not os.path.exists(config_file_path):
            raise FileNotFoundError(f"Configuration file '{config_file_path}' does not exist.")
        config.read(config_file_path)
        if 'database' not in config:
            raise KeyError("Missing 'database' section in the config file.")
        return config
    except Exception as e:
        logging.error(f"Error reading configuration file: {e}")
        sys.exit(1)

def get_db_pool(config):
    try:
        db_pool_args = {
            'user': config['database']['user'],
            'host': config['database']['host'],
            'port': config['database']['port'],
            'database': config['database']['dbname']
        }
        password = config['database'].get('password')
        if password:
            db_pool_args['password'] = password
            logging.info("Using password authentication.")
        else:
            logging.info("No password in config; relying on .pgpass or other method.")
        db_pool = pool.SimpleConnectionPool(1, 10, **db_pool_args)
        if db_pool:
            logging.info("Connection pool created successfully.")
        return db_pool
    except DatabaseError as e:
        logging.error(f"Database error creating pool: {e}")
        sys.exit(1)
    except KeyError as e:
        logging.error(f"Missing config key: {e}")
        sys.exit(1)
    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        sys.exit(1)

def softmax(scores):
    exp_scores = np.exp(scores - np.max(scores)) 
    return exp_scores / np.sum(exp_scores)

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    config = read_config(script_dir)
    log_file = "/home/exx/myCode/horse-racing/FoxRiverAIRacing/logs/Predictions.log"
    setup_logging(script_dir, log_file)

    db_pool = get_db_pool(config)
    spark, jdbc_url, jdbc_properties, parquet_dir, _ = initialize_environment()

    # Connect to Postgres
    conn = db_pool.getconn()

    # Example query
    sql = """
    SELECT
        course_cd,
        track_name,
        race_date,
        race_number,
        saddle_cloth_number,
        horse_name,
        global_speed_score,
        "YetiRank:top=4_NDCG:top=1" AS model1_score,
        "YetiRank:top=1_NDCG:top=3" AS model2_score
    FROM predictions_2025_02_18_1
    ORDER BY course_cd, race_date, race_number, saddle_cloth_number
    """

    df = pd.read_sql(sql, conn)
    conn.close()

    # Convert race_date
    df["race_date"] = pd.to_datetime(df["race_date"], errors="coerce")

    group_cols = ["course_cd", "race_date", "race_number"]
    grouped = df.groupby(group_cols, sort=False)

    # ---------------------------------------------------
    #  A) Print to console
    # ---------------------------------------------------
    for (course, rdate, rnum), group in grouped:
        print("="*60)
        print(f"Race: {course} | Track: {group.iloc[0]['track_name']} | Date: {rdate.strftime('%Y-%m-%d')} | Race #: {int(rnum)}")
        print("-"*60)
        group_sorted = group.sort_values("model1_score", ascending=False).copy()

        scores = group_sorted[["model1_score","model2_score"]].values
        probs = np.apply_along_axis(softmax, 1, scores)
        group_sorted["model1_prob"] = probs[:, 0]
        group_sorted["model2_prob"] = probs[:, 1]

        for _, row in group_sorted.iterrows():
            scloth = row["saddle_cloth_number"]
            hname = row["horse_name"]
            speedfig = row["global_speed_score"]
            m1_score = row["model1_score"]
            m2_score = row["model2_score"]
            m1_prob = row["model1_prob"]
            m2_prob = row["model2_prob"]
            print(f"Saddle#: {scloth:<3} Horse: {hname:<20} "
                  f"SpeedFig: {speedfig:6.1f}  "
                  f"M1_Score: {m1_score:8.2f} ({m1_prob*100:5.1f}%)  "
                  f"M2_Score: {m2_score:8.2f} ({m2_prob*100:5.1f}%)")

        print("="*60, "\n")

    # ---------------------------------------------------
    #  B) Write to a DOCX file
    # ---------------------------------------------------
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("CatBoost Predictions Summary", 0)

    # We'll loop again to produce the docx content
    for (course, rdate, rnum), group in grouped:
        doc.add_heading(f"Race: {course} | {rdate.strftime('%Y-%m-%d')} | Race#{int(rnum)}", level=1)
        doc.add_paragraph(f"Track: {group.iloc[0]['track_name']}")

        group_sorted = group.sort_values("model1_score", ascending=False).copy()
        scores = group_sorted[["model1_score","model2_score"]].values
        probs = np.apply_along_axis(softmax, 1, scores)
        group_sorted["model1_prob"] = probs[:, 0]
        group_sorted["model2_prob"] = probs[:, 1]

        # Create a table with 6 columns
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Saddle#"
        hdr_cells[1].text = "Horse Name"
        hdr_cells[2].text = "SpeedFig"
        hdr_cells[3].text = "M1 Score"
        hdr_cells[4].text = "M1 %"
        hdr_cells[5].text = "M2 Score/M2 %"

        for _, row in group_sorted.iterrows():
            scloth = row["saddle_cloth_number"]
            hname = row["horse_name"]
            speedfig = row["global_speed_score"]
            m1_score = row["model1_score"]
            m2_score = row["model2_score"]
            m1_prob = row["model1_prob"] * 100
            m2_prob = row["model2_prob"] * 100

            row_cells = table.add_row().cells
            row_cells[0].text = str(scloth)
            row_cells[1].text = str(hname)
            row_cells[2].text = f"{speedfig:.1f}"
            row_cells[3].text = f"{m1_score:.2f}"
            row_cells[4].text = f"{m1_prob:5.1f}%"
            row_cells[5].text = f"{m2_score:.2f}  ({m2_prob:5.1f}%)"

        doc.add_page_break()  # optional page break after each race

    # Save the DOCX
    output_docx = "/home/exx/myCode/horse-racing/FoxRiverAIRacing/data/predictions/predictions_summary.docx"
    doc.save(output_docx)
    logging.info(f"Saved predictions summary to {output_docx}")

if __name__ == "__main__":
    main()