import os
import sys
import logging
import psycopg2
import pandas as pd
import numpy as np
from datetime import datetime
import configparser
from docx import Document
from docx.shared import Pt

from psycopg2 import pool, DatabaseError
from src.data_preprocessing.data_prep1.data_utils import initialize_environment


############################################################
# 1) Setup logging and config
############################################################
def setup_logging(script_dir, log_file):
    try:
        with open(log_file, 'w'):
            pass
        logger = logging.getLogger()
        if logger.hasHandlers():
            logger.handlers.clear()
        logger.setLevel(logging.INFO)

        file_handler = logging.FileHandler(log_file)
        file_handler.setLevel(logging.INFO)
        formatter = logging.Formatter(
            '%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)

        logging.info("Logging initialized.")
    except Exception as e:
        print(f"Failed to set up logging: {e}", file=sys.stderr)
        sys.exit(1)

def read_config(script_dir, config_relative_path='../../config.ini'):
    try:
        config = configparser.ConfigParser()
        config_file_path = os.path.join(script_dir, config_relative_path)
        if not os.path.exists(config_file_path):
            raise FileNotFoundError(f"Configuration file '{config_file_path}' does not exist.")
        config.read(config_file_path)
        if 'database' not in config:
            raise KeyError("Missing 'database' section in the config file.")
        return config
    except Exception as e:
        logging.error(f"Error reading configuration file: {e}")
        sys.exit(1)

def get_db_pool(config):
    try:
        db_pool_args = {
            'user': config['database']['user'],
            'host': config['database']['host'],
            'port': config['database']['port'],
            'database': config['database']['dbname']
        }
        password = config['database'].get('password')
        if password:
            db_pool_args['password'] = password
            logging.info("Using password authentication.")
        else:
            logging.info("No password in config; relying on .pgpass or other method.")

        db_pool = pool.SimpleConnectionPool(1, 10, **db_pool_args)
        if db_pool:
            logging.info("Connection pool created successfully.")
        return db_pool
    except DatabaseError as e:
        logging.error(f"Database error creating pool: {e}")
        sys.exit(1)
    except KeyError as e:
        logging.error(f"Missing config key: {e}")
        sys.exit(1)
    except Exception as e:
        logging.error(f"Unexpected error: {e}")
        sys.exit(1)


############################################################
# 2) Softmax / Normalization
############################################################
def softmax(arr):
    """Compute a 1D softmax for an array of scores."""
    arr = np.array(arr, dtype=float)
    exp_scores = np.exp(arr - np.max(arr))
    return exp_scores / np.sum(exp_scores)


############################################################
# 3) Main Logic
############################################################
def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    config = read_config(script_dir)
    log_file = "/home/exx/myCode/horse-racing/FoxRiverAIRacing/logs/Predictions.log"
    setup_logging(script_dir, log_file)

    db_pool = get_db_pool(config)
    spark, jdbc_url, jdbc_properties, parquet_dir, _ = initialize_environment()

    conn = db_pool.getconn()

    # --------------------------------------------------------
    # A) DEFINE: The model columns with their top-4 success rates
    #
    # IMPORTANT:
    # The keys must match EXACTLY how they appear in the Postgres table,
    # which in your DDL is double-quoted.
    # We'll store them in a Python string with the double quotes included.
    # --------------------------------------------------------
    model_performance = {
        "yetirank3_ndcg4": 0.9535,
        "yetirank3_ndcg1": 0.9429,
        "yetirank2_ndcg1": 0.9360,
        "yetirank2_ndcg4": 0.9333,
        "yetirank1_ndcg1": 0.9285,
        "yetirank2_ndcg2": 0.9260,
        "queryrmse_ndcg3": 0.9256,
        "queryrmse_ndcg2": 0.9255,
        "yetirank1_ndcg2": 0.9195,
        "yetirank1_ndcg4": 0.9119,
        "yetirank3_ndcg2": 0.8991,
        "yetirank4_ndcg4": 0.8912,
        "yetirank2_ndcg3": 0.8906,
        "yetirank4_ndcg2": 0.8859,
        "yetirank1_ndcg3": 0.8835,
        "queryrmse_ndcg4": 0.8830,
        "yetirank4_ndcg3": 0.8751,
        "queryrmse_ndcg1": 0.8733,
        "yetirank4_ndcg1": 0.8446,
        "yetirank3_ndcg3": 0.8403,
    }

    # How much to blend ensemble vs. speed figure
    alpha_blend = 0.80

    # --------------------------------------------------------
    # B) Build the SQL 
    # We'll also retrieve missing_gps_flag, race_missing_flag
    # --------------------------------------------------------
    # The table name is 'predictions_2025_02_19_1'
    # Quoting the model columns EXACTLY as stored, plus we want missing_gps_flag, race_missing_flag
    # We'll do the same in Python so we can parse them as keys.
    columns_to_select = [
        "course_cd",
        "race_date",
        "race_number",
        "track_name",
        "saddle_cloth_number",
        "horse_name",
        "global_speed_score",
        "missing_gps_flag",
        "race_missing_flag",
    ] + list(model_performance.keys())  # our double-quoted model columns

    sql_select = ",\n       ".join(columns_to_select)
    sql = f"""
    SELECT
       {sql_select}
    FROM predictions_2025_02_19_1
    WHERE race_date >= CURRENT_DATE
    ORDER BY course_cd, race_date, race_number, saddle_cloth_number
    """

    df = pd.read_sql(sql, conn)  # May warn that psycopg2 is not fully tested, but it typically works.
    conn.close()

    # Convert date
    df["race_date"] = pd.to_datetime(df["race_date"], errors="coerce")

    # Group by race
    group_cols = ["course_cd", "race_date", "race_number"]
    grouped = df.groupby(group_cols, sort=False)

    # Convert the dictionary keys to a list for stable ordering
    model_cols = list(model_performance.keys())

    final_rows = []

    for (course, rdate, rnum), group in grouped:
        # 1) Build matrix => shape (num_horses, num_models)
        # These columns exist in df with the exact double-quoted names
        raw_scores_matrix = group[model_cols].values  # (nHorses, nModels)

        # 2) Per-model softmax => we do a column-based approach
        raw_scores_trans = raw_scores_matrix.T  # (nModels, nHorses)
        probs_trans = []
        for i, model_col in enumerate(model_cols):
            # one row of shape (nHorses,)
            row_scores = raw_scores_trans[i, :]
            row_probs = softmax(row_scores)
            probs_trans.append(row_probs)
        probs_trans = np.array(probs_trans)  # (nModels, nHorses)
        # transpose back => (nHorses, nModels)
        model_probs = probs_trans.T

        # 3) Weighted sum of these model probabilities
        weights = np.array([model_performance[m] for m in model_cols], dtype=float)
        # Weighted sum => shape (nHorses,)
        ensemble_scores = np.dot(model_probs, weights)
        # normalize across horses => sum=1
        sum_ens = np.sum(ensemble_scores)
        if sum_ens > 0:
            ensemble_scores /= sum_ens

        # 4) Optionally blend with speed figure distribution
        speedfig_array = group["global_speed_score"].fillna(0).values
        speedfig_probs = softmax(speedfig_array)
        final_probs = alpha_blend*ensemble_scores + (1.0 - alpha_blend)*speedfig_probs

        # Store results
        group_copy = group.copy()
        group_copy["ensemble_prob"] = final_probs

        # sort descending
        group_sorted = group_copy.sort_values("ensemble_prob", ascending=False)

        for idx, row in group_sorted.iterrows():
            final_rows.append({
                "course_cd": row["course_cd"],
                "race_date": row["race_date"],
                "race_number": row["race_number"],
                "track_name": row["track_name"],
                "saddle_cloth_number": row["saddle_cloth_number"],
                "horse_name": row["horse_name"],
                "global_speed_score": row["global_speed_score"],
                "missing_gps_flag": row["missing_gps_flag"],
                "race_missing_flag": row["race_missing_flag"],
                "ensemble_prob": row["ensemble_prob"]
            })

    final_df = pd.DataFrame(final_rows)
    final_df.sort_values(
        ["course_cd","race_date","race_number","ensemble_prob"],
        ascending=[True,True,True,False],
        inplace=True
    )

    # --------------------------------------------------------
    # Output to docx with missing data annotation
    # --------------------------------------------------------
    doc = Document()
    doc.add_heading("Ensemble Predictions (Weighted + Softmax)", 0)

    for (course, rdate, rnum), group in final_df.groupby(["course_cd","race_date","race_number"], sort=False):
        doc.add_heading(f"Race: {course} | {rdate.strftime('%Y-%m-%d')} | Race#{int(rnum)}", level=1)
        track_nm = group.iloc[0]["track_name"]
        doc.add_paragraph(f"Track: {track_nm}")

        group_sorted = group.sort_values("ensemble_prob", ascending=False)

        # Check if any horse is missing data
        any_missing = (group_sorted["missing_gps_flag"] == 1).any() or (group_sorted["race_missing_flag"] == 1).any()

        # Create table
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Saddle#"
        hdr_cells[1].text = "Horse Name"
        hdr_cells[2].text = "SpeedFig"
        hdr_cells[3].text = "Ensemble %"

        for _, row in group_sorted.iterrows():
            row_cells = table.add_row().cells
            # If missing, we add an asterisk next to name
            star = ""
            if row["missing_gps_flag"] == 1 or row["race_missing_flag"] == 1:
                star = "*"

            row_cells[0].text = str(row["saddle_cloth_number"])
            row_cells[1].text = f"{row['horse_name']}{star}"
            row_cells[2].text = f"{row['global_speed_score']:.1f}"
            row_cells[3].text = f"{row['ensemble_prob']*100:5.2f}%"

        # If any missing data => note
        if any_missing:
            doc.add_paragraph(
                "NOTE: THIS RACE HAS HORSES THAT ARE MISSING CRITICAL DATA ELEMENTS"
            )

        doc.add_page_break()

    out_path = "/home/exx/myCode/horse-racing/FoxRiverAIRacing/data/predictions/final_ensemble_predictions.docx"
    doc.save(out_path)
    logging.info(f"Saved predictions summary to {out_path}")

if __name__ == "__main__":
    main()